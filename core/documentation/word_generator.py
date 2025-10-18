"""Word report generation for Power BI model documentation."""

from __future__ import annotations

import os
from typing import Any, Dict, List, Optional, Tuple

from .complexity_analyzer import calculate_measure_complexity
from .utils import DEFAULT_BRANDING, ensure_dir, now_iso, safe_filename, truncate


def _add_bookmark(paragraph, bookmark_name: str) -> None:
    """Add a bookmark to a paragraph for cross-referencing."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Create bookmark start
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), str(hash(bookmark_name) % 10000))
        bookmark_start.set(qn("w:name"), bookmark_name)

        # Create bookmark end
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(hash(bookmark_name) % 10000))

        # Insert at paragraph level
        paragraph._element.insert(0, bookmark_start)
        paragraph._element.append(bookmark_end)
    except Exception:
        pass  # Bookmark failed, continue without it


def _add_hyperlink(
    paragraph,
    text: str,
    bookmark_name: str,
    color_rgb: Tuple[int, int, int] = (0, 102, 204),
) -> None:
    """Add a hyperlink to a bookmark within the document."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Create hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), bookmark_name)

        # Create run inside hyperlink
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Style as hyperlink (blue, underlined)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "%02x%02x%02x" % color_rgb)
        rPr.append(color)

        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

        run.append(rPr)

        # Add text
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        run.append(text_elem)

        hyperlink.append(run)
        paragraph._element.append(hyperlink)
    except Exception:
        # Fallback to regular text if hyperlink fails
        paragraph.add_run(text)


def convert_to_pdf(docx_path: str, pdf_path: str) -> Dict[str, Any]:
    """Convert Word document to PDF using available methods."""
    try:
        # Method 1: Try docx2pdf (cross-platform, requires LibreOffice or Word)
        try:
            from docx2pdf import convert

            convert(docx_path, pdf_path)
            return {"success": True, "pdf_path": pdf_path}
        except ImportError:
            pass

        # Method 2: Try comtypes (Windows only, requires Word)
        try:
            import comtypes.client

            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = PDF format
            doc.Close()
            word.Quit()
            return {"success": True, "pdf_path": pdf_path}
        except Exception:
            pass

        # Method 3: Try win32com (Windows only, alternative)
        try:
            import win32com.client

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close()
            word.Quit()
            return {"success": True, "pdf_path": pdf_path}
        except Exception:
            pass

        return {
            "success": False,
            "error": "No PDF conversion method available. Install docx2pdf, LibreOffice, or MS Word.",
        }
    except Exception as exc:
        return {"success": False, "error": str(exc)}


def _format_table_professional(table, has_header: bool = True) -> None:
    """Apply professional styling to a Word table with header row and zebra striping."""
    try:
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Apply table style
        table.style = "Light Grid Accent 1"

        if has_header and len(table.rows) > 0:
            # Style header row
            header_cells = table.rows[0].cells
            for cell in header_cells:
                # Header background color (dark blue)
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "00467F")  # Dark blue background
                cell._element.get_or_add_tcPr().append(shading_elm)

                # Header text formatting (white, bold)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(11)

        # Apply zebra striping to data rows
        for idx, row in enumerate(table.rows[1:], start=1):
            if idx % 2 == 0:
                # Even rows - light gray
                for cell in row.cells:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "F2F2F2")  # Light gray
                    cell._element.get_or_add_tcPr().append(shading_elm)

        # Set column widths and cell padding
        for row in table.rows:
            for cell in row.cells:
                # Add cell padding
                tc_pr = cell._element.get_or_add_tcPr()
                tc_mar = OxmlElement("w:tcMar")
                for margin_name in ["top", "left", "bottom", "right"]:
                    node = OxmlElement(f"w:{margin_name}")
                    node.set(qn("w:w"), "100")
                    node.set(qn("w:type"), "dxa")
                    tc_mar.append(node)
                tc_pr.append(tc_mar)
    except Exception:
        # If styling fails, table will still have content
        pass


def render_word_report(
    context: Dict[str, Any],
    *,
    output_dir: Optional[str] = None,
    graph_path: Optional[str] = None,
    graph_notes: Optional[List[str]] = None,
    change_summary: Optional[Dict[str, Any]] = None,
    mode: str = "full",
    branding: Optional[Dict[str, Any]] = None,
    add_hyperlinks: bool = True,
    export_pdf: bool = False,
) -> Dict[str, Any]:
    """Generate a comprehensive Word report for a Power BI model.

    Args:
        context: Documentation context dictionary from collect_model_documentation()
        output_dir: Optional output directory for the report
        graph_path: Optional path to relationship graph image
        graph_notes: Optional notes about graph generation
        change_summary: Optional change summary from snapshot diff
        mode: Report mode ("full" or "update")
        branding: Optional branding configuration dict
        add_hyperlinks: Whether to add internal hyperlinks (default True)
        export_pdf: Whether to also export as PDF (default False)

    Returns:
        dict: {
            "success": bool,
            "doc_path": str (path to Word document),
            "pdf_path": str (path to PDF, if exported),
            "pdf_warning": str (if PDF export failed)
        }
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except Exception as exc:
        return {"success": False, "error": f"python-docx not available: {exc}"}

    out_dir = ensure_dir(output_dir)
    doc_name = (
        safe_filename(
            context.get("database_name"), f"model_documentation_{now_iso()}"
        )
        + ".docx"
    )
    doc_path = os.path.join(out_dir, doc_name)

    doc = Document()

    # Apply custom branding or use defaults
    brand = {**DEFAULT_BRANDING, **(branding or {})}

    # Enhanced styling with professional theme (customizable)
    try:
        # Set default font and paragraph spacing
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Heading 1 - Primary color
        h1_style = doc.styles["Heading 1"]
        h1_style.font.name = "Calibri"
        h1_style.font.size = Pt(16)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(*brand["primary_color"])

        # Heading 2 - Secondary color
        h2_style = doc.styles["Heading 2"]
        h2_style.font.name = "Calibri"
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(*brand["secondary_color"])

        # Heading 3 - Accent color
        h3_style = doc.styles["Heading 3"]
        h3_style.font.name = "Calibri"
        h3_style.font.size = Pt(12)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(*brand["accent_color"])
    except Exception:
        pass

    # Add logo if provided
    if brand.get("logo_path") and os.path.exists(brand["logo_path"]):
        try:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(brand["logo_path"], width=Inches(2.0))
            doc.add_paragraph()  # Spacer
        except Exception:
            pass  # Logo failed to load, continue without it

    title = doc.add_heading(
        f"Power BI Model Documentation – {context.get('database_name', 'Power BI Model')}",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata section with styling
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"Generated on {context.get('generated_at')}")
    meta_run.font.size = Pt(10)
    meta_run.font.italic = True
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add company name if provided
    if brand.get("company_name"):
        company_para = doc.add_paragraph()
        company_run = company_para.add_run(brand["company_name"])
        company_run.font.size = Pt(9)
        company_run.font.italic = True
        company_run.font.color.rgb = RGBColor(128, 128, 128)
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if mode == "update":
        update_para = doc.add_paragraph()
        update_run = update_para.add_run(
            "Mode: Update – documentation refreshed with latest changes"
        )
        update_run.font.bold = True
        update_run.font.color.rgb = RGBColor(0, 112, 192)
        update_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add table of contents with hyperlinks
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    toc_para = doc.add_paragraph()
    toc_para.add_run("This document contains the following sections:\n")

    # Define sections with bookmark names
    sections = [
        ("Main Detailed Summary", "section_summary"),
        ("Tables", "section_tables"),
        ("Relationships", "section_relationships"),
        ("Dependencies", "section_dependencies"),
        ("Best Practice Analysis", "section_bpa"),
    ]
    if change_summary:
        sections.append(("Change Log", "section_changelog"))

    # Create TOC with hyperlinks if enabled
    for section_name, bookmark_name in sections:
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.paragraph_format.left_indent = Inches(0.25)
        if add_hyperlinks:
            _add_hyperlink(bullet, section_name, bookmark_name, brand["secondary_color"])
        else:
            bullet.add_run(section_name)

    doc.add_page_break()

    summary_heading = doc.add_heading("Main Detailed Summary", level=1)
    if add_hyperlinks:
        _add_bookmark(summary_heading, "section_summary")
    summary = context.get("summary", {})
    counts = summary.get("counts", {}) if isinstance(summary, dict) else {}
    summary_list = doc.add_paragraph()
    summary_list.style = "List Bullet"
    summary_list.add_run(f"Tables: {counts.get('tables', 'n/a')}")
    summary_list = doc.add_paragraph()
    summary_list.style = "List Bullet"
    summary_list.add_run(f"Columns: {counts.get('columns', 'n/a')}")
    summary_list = doc.add_paragraph()
    summary_list.style = "List Bullet"
    summary_list.add_run(f"Measures: {counts.get('measures', 'n/a')}")
    summary_list = doc.add_paragraph()
    summary_list.style = "List Bullet"
    summary_list.add_run(f"Relationships: {counts.get('relationships', 'n/a')}")

    narrative = context.get("narrative", {})
    if narrative and narrative.get("text"):
        doc.add_paragraph(narrative.get("text"))
        highlights = narrative.get("highlights") or []
        if highlights:
            doc.add_heading("Highlights", level=2)
            for item in highlights:
                para = doc.add_paragraph(item, style="List Bullet")
                para.paragraph_format.left_indent = Inches(0.25)

    # Tables section
    doc.add_page_break()
    tables_heading = doc.add_heading("Tables", level=1)
    if add_hyperlinks:
        _add_bookmark(tables_heading, "section_tables")
    tables = context.get("tables", [])
    table_summary = doc.add_table(rows=1, cols=5)
    hdr = table_summary.rows[0].cells
    hdr[0].text = "Table"
    hdr[1].text = "Description"
    hdr[2].text = "Columns"
    hdr[3].text = "Measures"
    hdr[4].text = "Hidden"
    for tbl in tables:
        row = table_summary.add_row().cells
        row[0].text = tbl.get("name", "")
        row[1].text = tbl.get("description") or ""
        row[2].text = str(len(tbl.get("columns", [])))
        row[3].text = str(len(tbl.get("measures", [])))
        row[4].text = "Yes" if tbl.get("hidden") else "No"
    _format_table_professional(table_summary, has_header=True)
    for tbl in tables:
        doc.add_heading(tbl.get("name", "Unnamed Table"), level=2)
        if tbl.get("description"):
            doc.add_paragraph(tbl.get("description"))
        if tbl.get("columns"):
            doc.add_paragraph("Columns:", style="Heading 3")
            for col in tbl.get("columns"):
                desc_parts = [f"{col.get('name', '')} ({col.get('data_type', '')})"]
                if col.get("description"):
                    desc_parts.append(col.get("description"))
                if col.get("type") and col.get("type").lower() == "calculated":
                    desc_parts.append("calculated column")
                if col.get("hidden"):
                    desc_parts.append("hidden")
                para = doc.add_paragraph(
                    " – ".join([p for p in desc_parts if p]), style="List Bullet"
                )
                para.paragraph_format.left_indent = Inches(0.25)
        if tbl.get("measures"):
            doc.add_paragraph("Measures:", style="Heading 3")
            for meas in tbl.get("measures"):
                bullet = doc.add_paragraph(style="List Number")

                # Add measure name FIRST in bold
                name_run = bullet.add_run(meas.get("name", "Unnamed Measure"))
                name_run.font.bold = True

                # Calculate and display complexity on the same line
                expr = meas.get("expression") or ""
                complexity = calculate_measure_complexity(
                    expr, meas.get("dependencies", {})
                )
                if complexity["score"] > 0:
                    complexity_run = bullet.add_run(
                        f" [{complexity['level']} Complexity: {complexity['score']}/100]"
                    )
                    complexity_run.font.size = Pt(9)
                    complexity_run.font.color.rgb = RGBColor(*complexity["color"])
                    complexity_run.font.bold = True

                # Add description on a new line if available
                if meas.get("description"):
                    bullet.add_run("\n   ")
                    desc_run = bullet.add_run(
                        f"Description: {meas.get('description')}"
                    )
                    desc_run.font.italic = True

                # Add dependencies in structured format on new lines
                dependents = meas.get("dependencies", {})
                cols = dependents.get("columns", [])
                meas_refs = dependents.get("measures", [])
                if cols or meas_refs:
                    bullet.add_run("\n   ")
                    dep_run = bullet.add_run("Dependencies:")
                    dep_run.font.italic = True
                    if cols:
                        bullet.add_run(
                            f"\n      • Columns: {', '.join(cols[:6])}{'…' if len(cols) > 6 else ''}"
                        )
                    if meas_refs:
                        bullet.add_run(
                            f"\n      • Measures: {', '.join(meas_refs[:6])}{'…' if len(meas_refs) > 6 else ''}"
                        )
                expr = meas.get("expression") or ""
                if expr:
                    # Improved DAX code formatting with gray background
                    code = doc.add_paragraph()
                    code_run = code.add_run(truncate(expr, 800))
                    code_run.font.name = "Consolas"
                    code_run.font.size = Pt(9)
                    # Add light gray background for code block
                    try:
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn

                        shading_elm = OxmlElement("w:shd")
                        shading_elm.set(qn("w:fill"), "F5F5F5")  # Light gray background
                        code.paragraph_format._element.get_or_add_pPr().append(
                            shading_elm
                        )
                        code.paragraph_format.left_indent = Inches(0.5)
                        code.paragraph_format.right_indent = Inches(0.25)
                        code.paragraph_format.space_before = Pt(6)
                        code.paragraph_format.space_after = Pt(6)
                    except Exception:
                        pass

    # Relationships
    doc.add_page_break()
    rels_heading = doc.add_heading("Relationships", level=1)
    if add_hyperlinks:
        _add_bookmark(rels_heading, "section_relationships")
    rels = context.get("relationships", [])
    rel_table = doc.add_table(rows=1, cols=6)
    hdr = rel_table.rows[0].cells
    hdr[0].text = "From"
    hdr[1].text = "To"
    hdr[2].text = "Active"
    hdr[3].text = "Cardinality"
    hdr[4].text = "Direction"
    hdr[5].text = "Description"
    for rel in rels:
        row = rel_table.add_row().cells
        row[0].text = f"{rel.get('from_table', '')}[{rel.get('from_column', '')}]"
        row[1].text = f"{rel.get('to_table', '')}[{rel.get('to_column', '')}]"
        row[2].text = "Yes" if rel.get("is_active") else "No"
        row[3].text = rel.get("cardinality", "")
        row[4].text = rel.get("direction", "")
        row[5].text = (
            "Active relationship"
            if rel.get("is_active")
            else "Inactive relationship"
        )
    _format_table_professional(rel_table, has_header=True)

    if graph_path:
        doc.add_heading("Visual Overview of Relationships", level=1)
        try:
            doc.add_picture(graph_path, width=Inches(6.5))
        except Exception as exc:
            doc.add_paragraph(f"Unable to embed graph image: {exc}")
    if graph_notes:
        for note in graph_notes:
            doc.add_paragraph(note, style="List Bullet")

    # Dependencies
    doc.add_page_break()
    deps_heading = doc.add_heading("Dependencies", level=1)
    if add_hyperlinks:
        _add_bookmark(deps_heading, "section_dependencies")
    edges = [
        edge for edge in context.get("dependency_edges", []) if edge.get("targets")
    ]
    if edges:
        dep_table = doc.add_table(rows=1, cols=3)
        hdr = dep_table.rows[0].cells
        hdr[0].text = "Source Measure"
        hdr[1].text = "Targets"
        hdr[2].text = "Type"
        for edge in edges:
            row = dep_table.add_row().cells
            row[0].text = edge.get("source", "")
            row[1].text = ", ".join(edge.get("targets", []))
            row[2].text = edge.get("target_type", "")
        _format_table_professional(dep_table, has_header=True)
    else:
        doc.add_paragraph("No dependencies detected")

    # Best practices
    doc.add_page_break()
    bpa_heading = doc.add_heading("Best Practice Analysis", level=1)
    if add_hyperlinks:
        _add_bookmark(bpa_heading, "section_bpa")
    bp = context.get("best_practices", {})
    doc.add_paragraph(f"Source: {bp.get('source') or 'Not available'}")
    doc.add_paragraph(f"Total issues: {bp.get('total_issues', 0)}")
    if bp.get("by_severity"):
        sev_table = doc.add_table(rows=1, cols=2)
        hdr = sev_table.rows[0].cells
        hdr[0].text = "Severity"
        hdr[1].text = "Count"
        for severity, count in sorted(
            bp.get("by_severity", {}).items(), key=lambda kv: kv[0]
        ):
            row = sev_table.add_row().cells
            row[0].text = severity
            row[1].text = str(count)
        _format_table_professional(sev_table, has_header=True)
    if bp.get("violations"):
        doc.add_paragraph("Detailed violations:", style="Heading 2")
        for item in bp.get("violations", [])[:200]:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(
                f"[{item.get('severity', 'Info')}] {item.get('rule_name', item.get('rule_id', 'Rule'))}: {item.get('description', '')}"
            )
            detail = []
            if item.get("object_type") and item.get("object_name"):
                detail.append(f"Object: {item['object_type']} {item['object_name']}")
            if item.get("table_name"):
                detail.append(f"Table: {item['table_name']}")
            if detail:
                doc.add_paragraph("; ".join(detail), style="List Bullet")
    if bp.get("notes"):
        for note in bp.get("notes"):
            doc.add_paragraph(note, style="List Bullet")

    if context.get("notes"):
        doc.add_heading("Generation Notes", level=1)
        for note in context.get("notes"):
            doc.add_paragraph(note, style="List Bullet")

    if change_summary:
        # Add page break before change log
        doc.add_page_break()
        changelog_heading = doc.add_heading("Change Log", level=1)
        if add_hyperlinks:
            _add_bookmark(changelog_heading, "section_changelog")

        # Visual indicator for changes
        changes_detected = change_summary.get("changes_detected")
        change_para = doc.add_paragraph()
        change_run = change_para.add_run(
            f"Changes detected: {'Yes' if changes_detected else 'No'}"
        )
        change_run.font.bold = True
        if changes_detected:
            change_run.font.color.rgb = RGBColor(192, 0, 0)  # Red for changes
        else:
            change_run.font.color.rgb = RGBColor(0, 128, 0)  # Green for no changes

        # Create summary table for changes
        has_changes = False
        change_rows = []
        for key in ("tables", "measures", "relationships"):
            section = change_summary.get(key, {})
            if section:
                for category in ("added", "removed", "updated"):
                    items = section.get(category)
                    if items:
                        has_changes = True
                        change_rows.append(
                            (
                                key.title(),
                                category.title(),
                                len(items),
                                ", ".join(items[:5])
                                + ("..." if len(items) > 5 else ""),
                            )
                        )

        if has_changes:
            change_table = doc.add_table(rows=1, cols=4)
            hdr = change_table.rows[0].cells
            hdr[0].text = "Object Type"
            hdr[1].text = "Change Type"
            hdr[2].text = "Count"
            hdr[3].text = "Items (first 5)"
            for obj_type, change_type, count, items in change_rows:
                row = change_table.add_row().cells
                row[0].text = obj_type
                row[1].text = change_type
                row[2].text = str(count)
                row[3].text = items
            _format_table_professional(change_table, has_header=True)

        if change_summary.get("best_practices"):
            bp_change = change_summary["best_practices"]
            doc.add_heading("Best Practice Changes", level=2)
            doc.add_paragraph(f"Previous issues: {bp_change.get('previous', 'N/A')}")
            doc.add_paragraph(f"Current issues: {bp_change.get('current', 'N/A')}")
            delta = (
                bp_change.get("current", 0) - bp_change.get("previous", 0)
                if isinstance(bp_change.get("current"), int)
                and isinstance(bp_change.get("previous"), int)
                else None
            )
            if delta is not None:
                delta_para = doc.add_paragraph()
                delta_run = delta_para.add_run(
                    f"Delta: {'+' if delta > 0 else ''}{delta}"
                )
                delta_run.font.bold = True
                if delta > 0:
                    delta_run.font.color.rgb = RGBColor(192, 0, 0)  # Red for increase
                elif delta < 0:
                    delta_run.font.color.rgb = RGBColor(0, 128, 0)  # Green for decrease

    doc.save(doc_path)

    result = {
        "success": True,
        "doc_path": doc_path,
    }

    # PDF export if requested
    if export_pdf:
        pdf_path = doc_path.replace(".docx", ".pdf")
        pdf_result = convert_to_pdf(doc_path, pdf_path)
        if pdf_result.get("success"):
            result["pdf_path"] = pdf_path
        else:
            result["pdf_warning"] = pdf_result.get("error", "PDF conversion failed")

    return result
